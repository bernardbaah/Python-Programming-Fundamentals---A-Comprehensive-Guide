"""
Elegant redesign — Python Programming Fundamentals (DOCX interior)
Bernard Baah / Filly Coder
Amazon KDP 8.5 × 11 in — no cover

Layout mirrors build_pdf_elegant.py:
  • 2 stock photos per chapter (first + middle of pool)
  • 1 chart (chNN_chart.png) at ~halfway
  • 2 tables per chapter: tbl1 → real grid (budget ≤15), tbl2 → ref strip
  • tables[2]  → numbered process panel
  • key_concepts → callout box at ~2/5
  • 4 case studies per chapter
  • All text_figures → cream/gold code blocks
  • End of chapter: Key Concepts, Practice Exercises, Chapter Summary
"""
import json, os, re
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.shared import OxmlElement as OE

# ── Palette ───────────────────────────────────────────────────────────────────
_NAVY   = (0x0a, 0x16, 0x28)
_SLATE  = (0x1e, 0x3a, 0x5f)
_GOLD   = (0xc9, 0xa8, 0x4c)
_LGRAY  = (0xf2, 0xf5, 0xfa)
_MGRAY  = (0xdc, 0xe3, 0xee)
_DGRAY  = (0x4a, 0x4a, 0x4a)
_CGRAY  = (0x88, 0x88, 0x88)
_WHITE  = (0xff, 0xff, 0xff)
_CREAM  = (0xfe, 0xf9, 0xed)   # code block header
_CODETEXT = (0x1a, 0x27, 0x44) # dark navy code text

def _rgb(*t): return RGBColor(*t)

# ── Helpers ───────────────────────────────────────────────────────────────────
def sanitize(t):
    t = t.replace('\x00', '')
    t = re.sub(r'[\x01-\x08\x0b\x0c\x0e-\x1f\x7f]', '', t)
    return t

def set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex.upper().lstrip('#'))
    tcPr.append(shd)

def set_para_shading(para, fill_hex):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex.upper().lstrip('#'))
    pPr.append(shd)

def set_para_border(para, side, color='c9a84c', sz='18', space='4'):
    pPr  = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    el = OxmlElement(f'w:{side}')
    el.set(qn('w:val'),   'single')
    el.set(qn('w:sz'),    sz)
    el.set(qn('w:space'), space)
    el.set(qn('w:color'), color)
    pBdr.append(el)

def add_run(para, text, bold=False, italic=False, size=None, color=None,
            font='Times New Roman'):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = font
    if size:  run.font.size = Pt(size)
    if color: run.font.color.rgb = _rgb(*color)
    return run

def add_empty(doc, height_pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(height_pt)
    return p

# ── Text cleaning (mirrors build_pdf_elegant.py) ──────────────────────────────
_ARTIFACT_LINES = {
    'python programming fundamentals',
    'a comprehensive guide',
    'bernard baah',
    'filly coder',
    'filly coder · ai future series',
}

# Matches "CHAPTER N" or "CHAPTER N:" as a standalone line (to drop from body)
_CHAPTER_HDR_RE = re.compile(r'^CHAPTER\s+\d+:?\s*$', re.IGNORECASE)

# Matches numbered-section headings like "1.1", "1.1 " or "1.1 What is Python?"
_SECTION_NUM_RE = re.compile(r'^\d+\.\d+(\s+\S.*)?$')

# Matches short all-caps section labels (5–60 chars, no lowercase)
_ALLCAPS_HDR_RE = re.compile(r'^[A-Z][A-Z\s\-:]{4,59}$')

def _is_section_break(line):
    """True if this line should force a new paragraph (section heading)."""
    return bool(_SECTION_NUM_RE.match(line) or _ALLCAPS_HDR_RE.match(line))

def clean_lines(lines):
    """Remove running-header/footer artifacts and bare page numbers."""
    out = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            out.append('')      # preserve paragraph boundaries
            continue
        low = stripped.lower()
        if low in _ARTIFACT_LINES:
            continue
        if _CHAPTER_HDR_RE.match(stripped):     # "CHAPTER N:" label
            continue
        if re.match(r'^\d{1,3}$', stripped):    # bare page numbers
            continue
        out.append(stripped)
    return out

def merge_into_paragraphs(text):
    """
    Merge word-wrapped source lines into logical paragraphs.
    Paragraph boundaries: blank lines OR section-heading lines.
    Keeps section headings as their own single-line paragraph.
    """
    raw_lines = clean_lines(text.split('\n'))
    paras, current = [], []
    for line in raw_lines:
        if line == '':
            # Explicit blank line → flush
            if current:
                paras.append(' '.join(current))
                current = []
        elif _is_section_break(line):
            # Section heading → flush current, emit heading as its own para
            if current:
                paras.append(' '.join(current))
                current = []
            paras.append(line)
        else:
            current.append(line)
    if current:
        paras.append(' '.join(current))
    return paras

# ── Load data ─────────────────────────────────────────────────────────────────
with open('book_data/chapters_text.json') as f:
    raw = json.load(f)
ch_text = raw['chapters']

ch_data = {}
for n in range(1, 26):
    path = f'book_data/chapter_{n:02d}.json'
    if os.path.exists(path):
        with open(path) as f:
            ch_data[n] = json.load(f)
print(f'Loaded {len(ch_data)} chapters')

# ── Document ──────────────────────────────────────────────────────────────────
doc = Document()

# Page setup
sec = doc.sections[0]
sec.page_width    = Inches(8.5)
sec.page_height   = Inches(11)
sec.left_margin   = Inches(1.0)
sec.right_margin  = Inches(1.0)
sec.top_margin    = Inches(1.1)
sec.bottom_margin = Inches(0.9)

# ── Header / footer helpers ───────────────────────────────────────────────────
# Text width 6.5 in → right-tab at 9360 twips (6.5 × 1440)
_TAB_RIGHT_TWIPS = '9360'
_BOOK_TITLE = 'PYTHON PROGRAMMING FUNDAMENTALS'

def _add_right_tab_stop(para):
    """Attach a right-aligned tab stop at the far end of the text width."""
    pPr  = para._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab  = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), _TAB_RIGHT_TWIPS)
    tabs.append(tab)
    # Remove any existing tabs element before appending
    existing = pPr.find(qn('w:tabs'))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(tabs)


def _insert_page_number(para):
    """Append a { PAGE } field to *para* using three OxmlElement runs."""
    def _run(child):
        r = OxmlElement('w:r')
        # Copy run-properties from the paragraph's default font
        rPr = OxmlElement('w:rPr')
        color_el = OxmlElement('w:color')
        color_el.set(qn('w:val'), '%02X%02X%02X' % _DGRAY)
        rPr.append(color_el)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '18')   # 9 pt × 2
        rPr.append(sz)
        fn = OxmlElement('w:rFonts')
        fn.set(qn('w:ascii'), 'Calibri')
        fn.set(qn('w:hAnsi'), 'Calibri')
        rPr.append(fn)
        r.append(rPr)
        r.append(child)
        return r

    begin = OxmlElement('w:fldChar');  begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = ' PAGE '
    end   = OxmlElement('w:fldChar');  end.set(qn('w:fldCharType'), 'end')

    para._p.append(_run(begin))
    para._p.append(_run(instr))
    para._p.append(_run(end))


def _setup_section_hf(section, chapter_title=''):
    """
    Configure header and footer for *section*.

    Header  — left:  PYTHON PROGRAMMING FUNDAMENTALS (navy, Calibri 8)
               right: chapter title (gold, Calibri 8)
               bottom: thin gold rule
    Footer  — centred PAGE field (Calibri 9, dark-gray)
               top:    thin gold rule
    """
    # ── Header ────────────────────────────────────────────────────────────────
    hdr = section.header
    hdr.is_linked_to_previous = False

    # python-docx always gives at least one paragraph in a fresh header
    hp = hdr.paragraphs[0]
    hp.clear()
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after  = Pt(3)

    _add_right_tab_stop(hp)

    r_left = hp.add_run(_BOOK_TITLE)
    r_left.font.name  = 'Calibri'
    r_left.font.size  = Pt(8)
    r_left.font.color.rgb = _rgb(*_NAVY)

    hp.add_run('\t')   # jump to right tab stop

    right_text = chapter_title.upper()[:60] if chapter_title else ''
    r_right = hp.add_run(right_text)
    r_right.font.name  = 'Calibri'
    r_right.font.size  = Pt(8)
    r_right.font.color.rgb = _rgb(*_GOLD)

    set_para_border(hp, 'bottom', 'c9a84c', '6', '4')

    # ── Footer ────────────────────────────────────────────────────────────────
    ftr = section.footer
    ftr.is_linked_to_previous = False

    fp = ftr.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(4)
    fp.paragraph_format.space_after  = Pt(0)

    set_para_border(fp, 'top', 'c9a84c', '6', '4')
    _insert_page_number(fp)


# Apply header/footer to the initial (only at this point) section.
# The chapter loop will create additional sections — one per chapter — and
# call _setup_section_hf on each so the right header shows the chapter title.
_setup_section_hf(sec, '')   # front matter: no chapter title yet

# ── Global default style ──────────────────────────────────────────────────────
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(11)
normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
normal.paragraph_format.line_spacing      = 1.35
normal.paragraph_format.space_after       = Pt(6)

h1 = doc.styles['Heading 1']
h1.font.name = 'Calibri'; h1.font.size = Pt(22); h1.font.bold = True
h1.font.color.rgb = _rgb(*_WHITE)

h2 = doc.styles['Heading 2']
h2.font.name = 'Calibri'; h2.font.size = Pt(13); h2.font.bold = True
h2.font.color.rgb = _rgb(*_NAVY)
h2.paragraph_format.space_before = Pt(14)
h2.paragraph_format.space_after  = Pt(2)

h3 = doc.styles['Heading 3']
h3.font.name = 'Calibri'; h3.font.size = Pt(11); h3.font.bold = True
h3.font.color.rgb = _rgb(*_SLATE)

# ── Table budget (mirrors PDF: max 15 real grid tables across all chapters) ────
MAX_BOOK_TABLES   = 15
_book_tables_used = 0

# ── Helper: chapter banner ────────────────────────────────────────────────────
def add_chapter_banner(doc, num, title):
    """Navy shaded banner that mirrors the PDF chapter opener."""
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, '0a1628')
    cell.paragraphs[0].clear()

    lbl = cell.paragraphs[0]
    lbl.paragraph_format.space_before = Pt(14)
    lbl.paragraph_format.space_after  = Pt(2)
    r = lbl.add_run(f'CHAPTER  {num}')
    r.font.name = 'Calibri'; r.font.size = Pt(10); r.bold = False
    r.font.color.rgb = _rgb(*_GOLD)

    ttl = cell.add_paragraph(title)
    ttl.paragraph_format.space_before = Pt(4)
    ttl.paragraph_format.space_after  = Pt(18)
    r2 = ttl.runs[0]
    r2.font.name = 'Calibri'; r2.font.size = Pt(22); r2.bold = True
    r2.font.color.rgb = _rgb(*_WHITE)

    doc.add_paragraph()

# ── Helper: body paragraph ────────────────────────────────────────────────────
def add_body(doc, text, first=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if not first:
        p.paragraph_format.first_line_indent = Inches(0.2)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'; run.font.size = Pt(11)
    return p

def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    set_para_border(p, 'bottom', 'c9a84c', '6', '4')
    return p

# ── Helper: image ─────────────────────────────────────────────────────────────
def add_image_safe(doc, path, caption, width=Inches(5.4)):
    if path and os.path.exists(path):
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.add_run().add_picture(path, width=width)
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(10)
            r = cap.add_run(caption)
            r.italic = True; r.font.size = Pt(9)
            r.font.color.rgb = _rgb(*_CGRAY)
            r.font.name = 'Calibri'
            return True
        except Exception as e:
            print(f'  img err {path}: {e}')
    return False

# ── Helper: real grid table ───────────────────────────────────────────────────
def add_table_docx(doc, caption, headers, rows):
    if not headers: return
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(10)
    cap.paragraph_format.space_after  = Pt(4)
    r = cap.add_run(caption)
    r.bold = True; r.italic = True; r.font.size = Pt(10)
    r.font.color.rgb = _rgb(*_NAVY); r.font.name = 'Calibri'

    ncols = len(headers)
    tbl   = doc.add_table(rows=1 + len(rows), cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    hc = tbl.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hc[i], '0a1628')
        p = hc[i].paragraphs[0]
        p.clear()
        run = p.add_run(str(h))
        run.bold = True; run.font.size = Pt(9)
        run.font.name = 'Calibri'
        run.font.color.rgb = _rgb(*_WHITE)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)

    for ri, row in enumerate(rows):
        rc = tbl.rows[ri + 1].cells
        bg = 'f2f5fa' if ri % 2 == 0 else 'ffffff'
        for ci, val in enumerate(list(row)[:ncols]):
            set_cell_shading(rc[ci], bg)
            p = rc[ci].paragraphs[0]
            p.clear()
            run = p.add_run(str(val))
            run.font.size = Pt(9); run.font.name = 'Times New Roman'
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)

    doc.add_paragraph()

# ── Helper: reference strip (no grid, gold bottom accent) ────────────────────
def add_ref_strip(doc, caption, headers, rows):
    """Borderless reference table — used for tbl2 slot and budget fallback."""
    if not rows: return
    if caption:
        cap = doc.add_paragraph()
        cap.paragraph_format.space_before = Pt(8)
        cap.paragraph_format.space_after  = Pt(3)
        r = cap.add_run(caption)
        r.italic = True; r.font.size = Pt(8)
        r.font.color.rgb = _rgb(*_NAVY); r.font.name = 'Calibri'

    ncols = max(len(headers) if headers else 0,
                max((len(r) for r in rows), default=1))
    ncols = min(ncols, 5)
    tbl   = doc.add_table(rows=(1 if headers else 0) + len(rows), cols=ncols)
    tbl.style = 'Table Grid'   # we'll remove borders via cell shading + XML hack
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remove all borders via tblPr
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl._tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tblBorders.append(b)
    # Remove existing tblBorders if any
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)

    row_offset = 0
    if headers:
        hc = tbl.rows[0].cells
        for i, h in enumerate(list(headers)[:ncols]):
            set_cell_shading(hc[i], 'f0f0f0')
            p = hc[i].paragraphs[0]; p.clear()
            run = p.add_run(str(h))
            run.font.size = Pt(8); run.font.name = 'Calibri'
            run.font.color.rgb = _rgb(0x66, 0x66, 0x66)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
        row_offset = 1

    for ri, row in enumerate(rows):
        rc  = tbl.rows[ri + row_offset].cells
        bg  = 'fef9ed' if ri % 2 == 0 else 'ffffff'
        cells = list(row)[:ncols]
        while len(cells) < ncols:
            cells.append('')
        for ci, val in enumerate(cells):
            set_cell_shading(rc[ci], bg)
            p = rc[ci].paragraphs[0]; p.clear()
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Calibri' if ci == 0 else 'Times New Roman'
            if ci == 0: run.bold = True
            run.font.color.rgb = _rgb(*_NAVY) if ci == 0 else _rgb(*_DGRAY)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)

    # Gold bottom border on last row via paragraph border
    last_row = tbl.rows[-1]
    for cell in last_row.cells:
        for p in cell.paragraphs:
            set_para_border(p, 'bottom', 'c9a84c', '8', '2')

    doc.add_paragraph()

# ── Helper: callout box (key concepts) ───────────────────────────────────────
def add_callout_box(doc, concepts, max_items=4):
    """Slate left-stripe callout — mirrors PDF make_callout_box."""
    if not concepts: return
    items = concepts[:max_items]

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Narrow slate stripe | wide content area
    tbl.columns[0].width = Inches(0.18)
    tbl.columns[1].width = Inches(6.32)

    stripe_cell  = tbl.cell(0, 0)
    content_cell = tbl.cell(0, 1)

    set_cell_shading(stripe_cell,  '1e3a5f')   # slate
    set_cell_shading(content_cell, 'edf2fa')   # light blue-gray

    for p in stripe_cell.paragraphs:
        p.clear()

    content_cell.paragraphs[0].clear()
    lbl = content_cell.paragraphs[0]
    lbl.paragraph_format.space_before = Pt(8)
    lbl.paragraph_format.space_after  = Pt(4)
    r = lbl.add_run('💡  KEY CONCEPTS AT A GLANCE')
    r.bold = True; r.font.size = Pt(7)
    r.font.color.rgb = _rgb(*_GOLD); r.font.name = 'Calibri'

    for item in items:
        term = item.get('term', '').strip()
        defn = item.get('definition', '').strip()
        if len(defn) > 180:
            defn = defn[:177] + '…'
        if term:
            tp = content_cell.add_paragraph(term)
            tp.paragraph_format.space_before = Pt(4)
            tp.paragraph_format.space_after  = Pt(1)
            r2 = tp.runs[0]
            r2.bold = True; r2.font.size = Pt(9.5)
            r2.font.color.rgb = _rgb(*_NAVY); r2.font.name = 'Calibri'
        if defn:
            dp = content_cell.add_paragraph(defn)
            dp.paragraph_format.space_before = Pt(0)
            dp.paragraph_format.space_after  = Pt(4)
            r3 = dp.runs[0]
            r3.font.size = Pt(8.8); r3.font.name = 'Calibri'
            r3.font.color.rgb = _rgb(0x33, 0x33, 0x33)

    # Bottom padding paragraph
    pad = content_cell.add_paragraph()
    pad.paragraph_format.space_before = Pt(0)
    pad.paragraph_format.space_after  = Pt(6)

    doc.add_paragraph()

# ── Helper: process panel ─────────────────────────────────────────────────────
def add_process_panel(doc, title, headers, rows):
    """Navy header + numbered steps — mirrors PDF make_process_panel."""
    if not rows: return

    # Navy header bar
    hdr_tbl = doc.add_table(rows=1, cols=1)
    hdr_tbl.style = 'Table Grid'
    hdr_cell = hdr_tbl.cell(0, 0)
    set_cell_shading(hdr_cell, '0a1628')
    hdr_cell.paragraphs[0].clear()
    hp = hdr_cell.paragraphs[0]
    hp.paragraph_format.space_before = Pt(8)
    hp.paragraph_format.space_after  = Pt(8)
    r = hp.add_run(f'  ⚙  HOW IT WORKS: {title.upper()}')
    r.bold = True; r.font.size = Pt(9)
    r.font.color.rgb = _rgb(*_WHITE); r.font.name = 'Calibri'

    # Numbered steps table
    body_tbl = doc.add_table(rows=len(rows), cols=2)
    body_tbl.style = 'Table Grid'
    body_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    body_tbl.columns[0].width = Inches(0.42)
    body_tbl.columns[1].width = Inches(6.08)

    for i, row in enumerate(rows):
        cells = [str(c).strip() for c in row if str(c).strip()]
        if not cells: continue

        step_title = cells[0][:90]
        step_desc  = '  ·  '.join(cells[1:])[:200] if len(cells) > 1 else ''

        bg = 'f5f7fb' if i % 2 == 0 else 'ffffff'
        num_cell  = body_tbl.rows[i].cells[0]
        body_cell = body_tbl.rows[i].cells[1]

        set_cell_shading(num_cell,  bg)
        set_cell_shading(body_cell, bg)

        num_cell.paragraphs[0].clear()
        np_ = num_cell.paragraphs[0]
        np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
        np_.paragraph_format.space_before = Pt(7)
        np_.paragraph_format.space_after  = Pt(7)
        nr = np_.add_run(str(i + 1))
        nr.bold = True; nr.font.size = Pt(14)
        nr.font.color.rgb = _rgb(*_GOLD); nr.font.name = 'Calibri'

        body_cell.paragraphs[0].clear()
        bp = body_cell.paragraphs[0]
        bp.paragraph_format.space_before = Pt(7)
        bp.paragraph_format.space_after  = Pt(2) if step_desc else Pt(7)
        br = bp.add_run(step_title)
        br.bold = True; br.font.size = Pt(9.5)
        br.font.color.rgb = _rgb(*_NAVY); br.font.name = 'Calibri'

        if step_desc:
            desc_p = body_cell.add_paragraph(step_desc)
            desc_p.paragraph_format.space_before = Pt(1)
            desc_p.paragraph_format.space_after  = Pt(7)
            dr = desc_p.runs[0]
            dr.font.size = Pt(8.8); dr.font.name = 'Calibri'
            dr.font.color.rgb = _rgb(0x44, 0x44, 0x44)

    # Gold bottom border on last row
    last_row = body_tbl.rows[-1]
    for cell in last_row.cells:
        for p in cell.paragraphs:
            set_para_border(p, 'bottom', 'c9a84c', '8', '2')

    doc.add_paragraph()

# ── Helper: case study ────────────────────────────────────────────────────────
def add_case_study(doc, title, body):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.columns[0].width = Inches(0.18)
    tbl.columns[1].width = Inches(6.32)

    gold_cell    = tbl.cell(0, 0)
    content_cell = tbl.cell(0, 1)
    set_cell_shading(gold_cell,    'c9a84c')
    set_cell_shading(content_cell, 'f2f5fa')

    for p in gold_cell.paragraphs:
        p.clear()

    content_cell.paragraphs[0].clear()
    lbl = content_cell.paragraphs[0]
    lbl.paragraph_format.space_before = Pt(8)
    lbl.paragraph_format.space_after  = Pt(2)
    r = lbl.add_run('CASE STUDY')
    r.bold = True; r.font.size = Pt(8)
    r.font.color.rgb = _rgb(*_GOLD); r.font.name = 'Calibri'

    ttl = content_cell.add_paragraph(title)
    ttl.paragraph_format.space_before = Pt(0)
    ttl.paragraph_format.space_after  = Pt(6)
    r2 = ttl.runs[0]
    r2.bold = True; r2.font.size = Pt(11)
    r2.font.color.rgb = _rgb(*_NAVY); r2.font.name = 'Calibri'

    bod = content_cell.add_paragraph(body)
    bod.paragraph_format.space_before = Pt(0)
    bod.paragraph_format.space_after  = Pt(10)
    r3 = bod.runs[0]
    r3.font.size = Pt(10); r3.font.name = 'Times New Roman'
    r3.font.color.rgb = _rgb(*_DGRAY)

    doc.add_paragraph()

# ── Helper: code block (cream/gold — mirrors PDF make_figure) ─────────────────
def add_figure_box(doc, title, content):
    """
    Cream header + white body code block with gold accents.
    Mirrors the PDF's make_figure() design.
    """
    # Outer table: single column, cream header row + white body
    tbl = doc.add_table(rows=2, cols=1)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # ── Header row: cream background ──────────────────────────────────────────
    hdr_cell = tbl.cell(0, 0)
    set_cell_shading(hdr_cell, 'fef9ed')   # warm cream
    hdr_cell.paragraphs[0].clear()

    # "● ● ●  title  Python" header line
    hp = hdr_cell.paragraphs[0]
    hp.paragraph_format.space_before = Pt(8)
    hp.paragraph_format.space_after  = Pt(8)
    r_dots = hp.add_run('● ● ●  ')
    r_dots.font.size = Pt(8); r_dots.font.name = 'Calibri'
    r_dots.font.color.rgb = _rgb(0xc9, 0xa8, 0x4c)
    r_title = hp.add_run(title[:70])
    r_title.bold = True; r_title.font.size = Pt(9)
    r_title.font.color.rgb = _rgb(*_NAVY); r_title.font.name = 'Calibri'
    r_lang = hp.add_run('  Python')
    r_lang.italic = True; r_lang.font.size = Pt(7.5)
    r_lang.font.color.rgb = _rgb(0x9b, 0x7e, 0x3a); r_lang.font.name = 'Calibri'

    # ── Body row: white background ────────────────────────────────────────────
    body_cell = tbl.cell(1, 0)
    set_cell_shading(body_cell, 'ffffff')
    body_cell.paragraphs[0].clear()

    lines = content.split('\n')
    first = True
    for line in lines:
        if first:
            p = body_cell.paragraphs[0]
            first = False
        else:
            p = body_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1) if line.strip() else Pt(3)
        if line.strip():
            r = p.add_run(line)
            r.font.size = Pt(9); r.font.name = 'Courier New'
            r.font.color.rgb = _rgb(*_CODETEXT)
        # empty lines left blank

    # Padding at bottom of body
    pad = body_cell.add_paragraph()
    pad.paragraph_format.space_before = Pt(4)
    pad.paragraph_format.space_after  = Pt(4)

    # Gold top border on header row
    for p in hdr_cell.paragraphs:
        set_para_border(p, 'top',    'c9a84c', '18', '2')
        set_para_border(p, 'left',   'c9a84c', '6',  '2')
        set_para_border(p, 'right',  'c9a84c', '6',  '2')
        break

    doc.add_paragraph()

# ── Title page ────────────────────────────────────────────────────────────────
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(80)
tp.paragraph_format.space_after  = Pt(6)
r = tp.add_run('Python Programming Fundamentals')
r.font.name = 'Calibri'; r.font.size = Pt(28); r.bold = True
r.font.color.rgb = _rgb(*_NAVY)

sp = doc.add_paragraph('A Comprehensive Guide')
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp.runs[0].font.name = 'Calibri'; sp.runs[0].font.size = Pt(15)
sp.runs[0].font.color.rgb = _rgb(*_SLATE)
sp2 = doc.add_paragraph('Second Edition')
sp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp2.runs[0].font.name = 'Calibri'; sp2.runs[0].font.size = Pt(13)
sp2.runs[0].font.color.rgb = _rgb(*_GOLD)
sp2.runs[0].italic = True

doc.add_paragraph()
ap = doc.add_paragraph('Bernard Baah')
ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
ap.runs[0].font.name = 'Calibri'; ap.runs[0].font.size = Pt(13)
ap.runs[0].bold = True

fp = doc.add_paragraph('A Comprehensive Guide to Python Programming')
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.runs[0].font.name = 'Calibri'; fp.runs[0].font.size = Pt(11)
fp.runs[0].font.color.rgb = _rgb(*_GOLD)

doc.add_page_break()

# ── Table of Contents ─────────────────────────────────────────────────────────
toc_h = doc.add_heading('Table of Contents', level=1)
toc_h.runs[0].font.color.rgb = _rgb(*_NAVY)
toc_h.runs[0].font.name = 'Calibri'

CH_TITLES = {
    1:'Introduction to Python', 2:'Getting Started',
    3:'Variables and Data Types', 4:'Control Structures',
    5:'Functions', 6:'Modules and Libraries',
    7:'File Handling', 8:'Exception Handling',
    9:'Data Structures', 10:'Object-Oriented Programming: Basics',
    11:'Object-Oriented Programming: Advanced Concepts',
    12:'Error Handling and Debugging',
    13:'Working with Files and Directories (Part 2)',
    14:'Introduction to Testing',
    15:'Introduction to Modules and Packages',
    16:'Advanced Modules and Packages',
    17:'Working with External APIs',
    18:'Introduction to Web Development with Flask',
    19:'Intermediate Flask Development',
    20:'Introduction to Data Visualization with Matplotlib',
    21:'Advanced Data Visualization with Seaborn',
    22:'Introduction to Machine Learning with scikit-learn',
    23:'Web Scraping with Python',
    24:'Python Automation',
    25:'Conclusion and Next Steps',
}

PARTS = {1:'Part I: Python Foundations', 6:'Part II: Intermediate Python',
         13:'Part III: Advanced Topics', 17:'Part IV: Web, Data & Machine Learning'}

for n in range(1, 26):
    if n in PARTS:
        pp = doc.add_paragraph()
        pp.paragraph_format.space_before = Pt(8)
        r = pp.add_run(PARTS[n])
        r.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = _rgb(*_GOLD); r.font.name = 'Calibri'
    title = ch_data[n]['title'] if n in ch_data else CH_TITLES.get(n, f'Chapter {n}')
    ep = doc.add_paragraph(f'Chapter {n}  —  {title}')
    ep.runs[0].font.size = Pt(11)
    ep.runs[0].font.name = 'Times New Roman'

doc.add_page_break()

# ── Chapters ───────────────────────────────────────────────────────────────────
def _split_para(para, max_chars):
    """Break a single oversized paragraph at sentence boundaries."""
    if len(para) <= max_chars:
        return [para]
    parts, buf = [], []
    for sent in re.split(r'(?<=[.!?])\s+', para):
        if buf and sum(len(s) for s in buf) + len(sent) > max_chars:
            parts.append(' '.join(buf))
            buf = []
        buf.append(sent)
    if buf:
        parts.append(' '.join(buf))
    return parts if parts else [para]

def split_text(text, max_chars=1600):
    """Split merged logical paragraphs into render chunks."""
    raw_paras = merge_into_paragraphs(text)
    # Further split any paragraph that exceeds max_chars
    paras = []
    for p in raw_paras:
        paras.extend(_split_para(p, max_chars))
    chunks, cur, clen = [], [], 0
    for p in paras:
        if clen + len(p) > max_chars and cur:
            chunks.append(cur); cur, clen = [], 0
        cur.append(p); clen += len(p)
    if cur: chunks.append(cur)
    return chunks

for ch_num in range(1, 26):
    ch       = ch_data.get(ch_num, {})
    title    = ch.get('title', CH_TITLES.get(ch_num, f'Chapter {ch_num}'))
    images   = ch.get('images', [])
    tables   = ch.get('tables', [])
    css      = ch.get('case_studies', [])
    figs     = ch.get('text_figures', [])
    kc       = ch.get('key_concepts', [])
    exercises= ch.get('exercises', [])
    summary  = ch.get('chapter_summary', '')
    orig     = sanitize(ch_text.get(str(ch_num), ''))

    print(f'  Ch {ch_num}: {title}')

    # ── Curated 2 photos per chapter (first + middle) ────────────────────────
    photo_pool = [img for img in images if isinstance(img, dict)]
    photos_2   = []
    if photo_pool:
        photos_2.append(photo_pool[0])
    if len(photo_pool) > 1:
        photos_2.append(photo_pool[len(photo_pool) // 2])

    # ── Chart ────────────────────────────────────────────────────────────────
    chart_path = f'book_data/images/charts/ch{ch_num:02d}_chart.png'
    has_chart  = os.path.exists(chart_path)

    # ── Table layout mirrors PDF ──────────────────────────────────────────────
    tables_2 = tables[:2]                              # max 2 real/ref tables
    proc_src = tables[2] if len(tables) > 2 else None  # → process panel

    # Each chapter gets its own section so the right header shows the chapter title.
    # For ch 1 the initial section (already set up) is reused and updated.
    # For ch 2-25 a new section (= page break) is added; the previous chapter's
    # doc.add_page_break() is removed since the section break provides it.
    if ch_num == 1:
        _setup_section_hf(doc.sections[0], title)
    else:
        _new_sec = doc.add_section(WD_SECTION.NEW_PAGE)
        _setup_section_hf(_new_sec, title)

    add_chapter_banner(doc, ch_num, title)

    chunks   = split_text(orig, max_chars=1600)
    n_chunks = max(len(chunks), 1)

    # ── Placement triggers (chunk indices, 0-based) ───────────────────────────
    tbl1_at    = max(0, n_chunks // 5 - 1)
    photo1_at  = max(tbl1_at + 1,    n_chunks // 3 - 1)
    callout_at = max(photo1_at + 1,  n_chunks * 2 // 5)
    chart_at   = max(callout_at + 1, n_chunks // 2)
    tbl2_at    = max(chart_at + 1,   n_chunks * 3 // 5)
    photo2_at  = max(tbl2_at + 1,    n_chunks * 2 // 3)
    process_at = max(photo2_at + 1,  n_chunks * 4 // 5)
    cs_every   = max(1, n_chunks // max(len(css), 1))

    min_img_gap = max(2, n_chunks // 6)
    last_img_ci = -9999

    photo_placed   = [False, False]
    chart_placed   = False
    tbl1_placed    = False
    tbl2_placed    = False
    callout_placed = False
    process_placed = False
    cs_placed      = set()
    first_para     = True

    for ci, chunk in enumerate(chunks):
        for line in chunk:
            line = line.strip()
            if not line: continue
            elif re.match(r'^\d+\.\d+\s+\S', line) or \
                 (len(line) < 65 and line.upper() == line and len(line) > 5):
                add_h2(doc, line)
                first_para = True
            elif line.startswith(('▪', '•', '-', '–')):
                bp = doc.add_paragraph(style='List Bullet')
                bp.paragraph_format.space_after = Pt(3)
                r = bp.add_run(line.lstrip('▪•-– ').strip())
                r.font.size = Pt(11); r.font.name = 'Times New Roman'
            else:
                add_body(doc, line, first=first_para)
                first_para = False

        # ── Table 1 at ~1/5 ──────────────────────────────────────────────────
        if not tbl1_placed and ci >= tbl1_at and tables_2:
            t = tables_2[0]
            if _book_tables_used < MAX_BOOK_TABLES:
                try:
                    add_table_docx(doc, t.get('caption',''),
                                   t.get('headers',[]), t.get('rows',[]))
                    _book_tables_used += 1
                except Exception as e:
                    print(f'    tbl1 err: {e}')
                    add_ref_strip(doc, t.get('caption',''),
                                  t.get('headers',[]), t.get('rows',[]))
            else:
                add_ref_strip(doc, t.get('caption',''),
                              t.get('headers',[]), t.get('rows',[]))
            tbl1_placed = True

        # ── Photo 1 at ~1/3 ──────────────────────────────────────────────────
        if (not photo_placed[0] and ci >= photo1_at
                and photos_2 and ci - last_img_ci >= min_img_gap):
            img = photos_2[0]
            if add_image_safe(doc, img.get('path',''),
                              img.get('caption', f'Figure {ch_num}.1')):
                photo_placed[0] = True
                last_img_ci = ci

        # ── Callout box at ~2/5 ───────────────────────────────────────────────
        if not callout_placed and ci >= callout_at and kc:
            add_callout_box(doc, kc, max_items=4)
            callout_placed = True

        # ── Chart at ~1/2 ─────────────────────────────────────────────────────
        if (not chart_placed and has_chart and ci >= chart_at
                and ci - last_img_ci >= min_img_gap):
            caption = f'Figure {ch_num}.C: {title} — Key Concepts Visualised'
            if add_image_safe(doc, chart_path, caption, width=Inches(5.8)):
                chart_placed = True
                last_img_ci = ci

        # ── Ref strip at ~3/5 (tbl2 slot — always ref strip, no grid) ─────────
        if not tbl2_placed and ci >= tbl2_at and len(tables_2) > 1:
            t = tables_2[1]
            add_ref_strip(doc, t.get('caption',''),
                          t.get('headers',[]), t.get('rows',[]))
            tbl2_placed = True

        # ── Photo 2 at ~2/3 ──────────────────────────────────────────────────
        if (not photo_placed[1] and ci >= photo2_at
                and len(photos_2) > 1 and ci - last_img_ci >= min_img_gap):
            img = photos_2[1]
            if add_image_safe(doc, img.get('path',''),
                              img.get('caption', f'Figure {ch_num}.2')):
                photo_placed[1] = True
                last_img_ci = ci

        # ── Process panel at ~4/5 ─────────────────────────────────────────────
        if not process_placed and ci >= process_at and proc_src:
            add_process_panel(doc, proc_src.get('caption', title),
                              proc_src.get('headers', []),
                              proc_src.get('rows', []))
            process_placed = True

        # ── Case study: one every cs_every chunks ─────────────────────────────
        cs_slot = (ci + 1) // cs_every - 1
        if css and (ci + 1) % cs_every == 0 and 0 <= cs_slot < len(css) \
                and cs_slot not in cs_placed:
            add_case_study(doc, css[cs_slot].get('title',''),
                           css[cs_slot].get('body',''))
            cs_placed.add(cs_slot)

    # ── Flush anything not placed (very short chapters) ───────────────────────
    if not tbl1_placed and tables_2:
        t = tables_2[0]
        if _book_tables_used < MAX_BOOK_TABLES:
            try:
                add_table_docx(doc, t.get('caption',''),
                               t.get('headers',[]), t.get('rows',[]))
                _book_tables_used += 1
            except:
                add_ref_strip(doc, t.get('caption',''),
                              t.get('headers',[]), t.get('rows',[]))
        else:
            add_ref_strip(doc, t.get('caption',''),
                          t.get('headers',[]), t.get('rows',[]))

    if not callout_placed and kc:
        add_callout_box(doc, kc, max_items=4)

    if not tbl2_placed and len(tables_2) > 1:
        t = tables_2[1]
        add_ref_strip(doc, t.get('caption',''),
                      t.get('headers',[]), t.get('rows',[]))

    if not process_placed and proc_src:
        add_process_panel(doc, proc_src.get('caption', title),
                          proc_src.get('headers', []),
                          proc_src.get('rows', []))

    for ci2, cs in enumerate(css):
        if ci2 not in cs_placed:
            add_case_study(doc, cs.get('title',''), cs.get('body',''))

    if not photo_placed[0] and photos_2:
        add_image_safe(doc, photos_2[0].get('path',''),
                       photos_2[0].get('caption', f'Figure {ch_num}.1'))
    if not chart_placed and has_chart:
        add_image_safe(doc, chart_path,
                       f'Figure {ch_num}.C: {title} — Key Concepts Visualised',
                       width=Inches(5.8))
    if not photo_placed[1] and len(photos_2) > 1:
        add_image_safe(doc, photos_2[1].get('path',''),
                       photos_2[1].get('caption', f'Figure {ch_num}.2'))

    # ── Code blocks ───────────────────────────────────────────────────────────
    for fig in figs:
        add_figure_box(doc, fig.get('title', 'Figure'), fig.get('content', ''))

    # ── Key Concepts section ──────────────────────────────────────────────────
    if kc:
        p = doc.add_heading('Key Concepts', level=2)
        set_para_border(p, 'bottom', 'c9a84c', '6', '4')
        for item in kc:
            term = item.get('term', '').strip()
            defn = item.get('definition', '').strip()
            if term:
                tp = doc.add_paragraph()
                tp.paragraph_format.space_before = Pt(4)
                tp.paragraph_format.space_after  = Pt(1)
                r = tp.add_run(term)
                r.bold = True; r.font.size = Pt(10)
                r.font.color.rgb = _rgb(*_SLATE); r.font.name = 'Calibri'
            if defn:
                dp = doc.add_paragraph(defn)
                dp.paragraph_format.space_before = Pt(0)
                dp.paragraph_format.space_after  = Pt(5)
                dp.paragraph_format.left_indent  = Inches(0.15)
                r2 = dp.runs[0]
                r2.font.size = Pt(10); r2.font.name = 'Times New Roman'
                r2.font.color.rgb = _rgb(*_DGRAY)

    # ── Practice Exercises ────────────────────────────────────────────────────
    if exercises:
        p = doc.add_heading('Practice Exercises', level=2)
        set_para_border(p, 'bottom', 'c9a84c', '6', '4')
        for i, ex in enumerate(exercises, 1):
            diff = ex.get('difficulty', '').strip()
            q    = ex.get('question', '').strip()
            hint = ex.get('hint', '').strip()
            if q:
                qp = doc.add_paragraph()
                qp.paragraph_format.space_before = Pt(6)
                qp.paragraph_format.space_after  = Pt(2)
                r = qp.add_run(f'{i}.  {q}')
                r.bold = True; r.font.size = Pt(10)
                r.font.color.rgb = _rgb(*_SLATE); r.font.name = 'Calibri'
            if diff:
                dp = doc.add_paragraph(f'Difficulty: {diff}')
                dp.paragraph_format.space_before = Pt(0)
                dp.paragraph_format.space_after  = Pt(2)
                r2 = dp.runs[0]
                r2.font.size = Pt(8); r2.bold = True
                r2.font.color.rgb = _rgb(*_GOLD); r2.font.name = 'Calibri'
            if hint:
                hp = doc.add_paragraph(f'Hint: {hint}')
                hp.paragraph_format.space_before = Pt(0)
                hp.paragraph_format.space_after  = Pt(4)
                hp.paragraph_format.left_indent  = Inches(0.15)
                r3 = hp.runs[0]
                r3.italic = True; r3.font.size = Pt(9)
                r3.font.color.rgb = _rgb(*_CGRAY); r3.font.name = 'Times New Roman'

    # ── Chapter Summary ───────────────────────────────────────────────────────
    if summary:
        p = doc.add_heading('Chapter Summary', level=2)
        set_para_border(p, 'bottom', 'c9a84c', '6', '4')
        for para in (summary.split('\n\n') if '\n\n' in summary else [summary]):
            para = para.strip()
            if para:
                sp = doc.add_paragraph(para)
                sp.paragraph_format.space_before = Pt(4)
                sp.paragraph_format.space_after  = Pt(6)
                r = sp.runs[0]
                r.italic = True; r.font.size = Pt(11)
                r.font.color.rgb = _rgb(*_DGRAY); r.font.name = 'Times New Roman'

    # No explicit page break here: for chapters 1–24, the next chapter's
    # doc.add_section(WD_SECTION.NEW_PAGE) provides the page break.
    # Chapter 25 is the last chapter; the document simply ends.

OUT = 'book_data/Python_Fundamentals_Interior.docx'
doc.save(OUT)
sz = os.path.getsize(OUT) / 1e6
print(f'\n✓ DOCX → {OUT}  ({sz:.1f} MB)')

# ── Auto-sync to GitHub ───────────────────────────────────────────────────────
import sys as _sys
_sys.path.insert(0, os.path.join(os.path.dirname(__file__)))
try:
    from sync_to_github import sync as _sync_gh
    _sync_gh()
except Exception as _e:
    print(f'\n⚠  GitHub sync skipped: {_e}')
