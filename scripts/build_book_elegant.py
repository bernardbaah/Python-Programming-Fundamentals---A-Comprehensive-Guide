"""
Elegant redesign — Python Programming Fundamentals (DOCX interior)
Bernard Baah / Filly Coder
Amazon KDP 8.5 × 11 in — no cover
"""
import json, os, re
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.shared import OxmlElement as OE

# ── Helpers ───────────────────────────────────────────────────────────────────
def sanitize(t):
    t = t.replace('\x00', '')
    t = re.sub(r'[\x01-\x08\x0b\x0c\x0e-\x1f\x7f]', '', t)
    return t

def set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex.upper().lstrip('#'))
    tcPr.append(shd)

def set_para_shading(para, fill_hex):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex.upper().lstrip('#'))
    pPr.append(shd)

def set_para_border(para, side, color='c9a84c', sz='18', space='4'):
    pPr  = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    el = OxmlElement(f'w:{side}')
    el.set(qn('w:val'),   'single')
    el.set(qn('w:sz'),    sz)
    el.set(qn('w:space'), space)
    el.set(qn('w:color'), color)
    pBdr.append(el)

def add_run(para, text, bold=False, italic=False, size=None, color=None,
            font='Times New Roman', space_after=None):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = font
    if size:  run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*bytes.fromhex(color.lstrip('#')))
    return run

def add_empty(doc, height_pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(height_pt)
    return p

# ── Load data ─────────────────────────────────────────────────────────────────
with open('book_data/chapters_text.json') as f:
    raw = json.load(f)
ch_text = raw['chapters']

ch_data = {}
for n in range(1, 26):
    path = f'book_data/chapter_{n:02d}.json'
    if os.path.exists(path):
        with open(path) as f:
            ch_data[n] = json.load(f)
print(f'Loaded {len(ch_data)} chapters')

# ── Document ──────────────────────────────────────────────────────────────────
doc = Document()

# Page setup
sec = doc.sections[0]
sec.page_width   = Inches(8.5)
sec.page_height  = Inches(11)
sec.left_margin  = Inches(1.0)
sec.right_margin = Inches(1.0)
sec.top_margin   = Inches(1.1)
sec.bottom_margin= Inches(0.9)

# ── Global default style ──────────────────────────────────────────────────────
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(11)
normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
normal.paragraph_format.line_spacing      = 1.35
normal.paragraph_format.space_after       = Pt(6)

# Heading 1 → chapter title (we override per chapter)
h1 = doc.styles['Heading 1']
h1.font.name  = 'Calibri'
h1.font.size  = Pt(22)
h1.font.bold  = True
h1.font.color.rgb = RGBColor(0xff, 0xff, 0xff)

# Heading 2
h2 = doc.styles['Heading 2']
h2.font.name  = 'Calibri'
h2.font.size  = Pt(13)
h2.font.bold  = True
h2.font.color.rgb = RGBColor(0x0a, 0x16, 0x28)
h2.paragraph_format.space_before = Pt(14)
h2.paragraph_format.space_after  = Pt(2)

# Heading 3
h3 = doc.styles['Heading 3']
h3.font.name  = 'Calibri'
h3.font.size  = Pt(11)
h3.font.bold  = True
h3.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

# ── Helper: chapter banner ────────────────────────────────────────────────────
def add_chapter_banner(doc, num, title):
    """Navy shaded table that mimics the PDF chapter banner."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, '0a1628')

    # Remove default paragraph from cell
    cell.paragraphs[0].clear()

    # "CHAPTER N" in gold
    lbl = cell.paragraphs[0]
    lbl.paragraph_format.space_before = Pt(14)
    lbl.paragraph_format.space_after  = Pt(2)
    r = lbl.add_run(f'CHAPTER  {num}')
    r.font.name   = 'Calibri'
    r.font.size   = Pt(10)
    r.font.bold   = False
    r.font.color.rgb = RGBColor(0xc9, 0xa8, 0x4c)

    # Title in white
    ttl = cell.add_paragraph(title)
    ttl.paragraph_format.space_before = Pt(4)
    ttl.paragraph_format.space_after  = Pt(18)
    r2 = ttl.runs[0]
    r2.font.name  = 'Calibri'
    r2.font.size  = Pt(22)
    r2.font.bold  = True
    r2.font.color.rgb = RGBColor(0xff, 0xff, 0xff)

    doc.add_paragraph()  # spacer

# ── Helper: add body paragraph ────────────────────────────────────────────────
def add_body(doc, text, first=False, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if indent and not first:
        p.paragraph_format.first_line_indent = Inches(0.2)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    return p

def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    # Add gold bottom border
    set_para_border(p, 'bottom', 'c9a84c', '6', '4')
    return p

# ── Helper: add image ─────────────────────────────────────────────────────────
def add_image_safe(doc, path, caption):
    if path and os.path.exists(path):
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            run = p.add_run()
            run.add_picture(path, width=Inches(5.4))
            # Caption
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(10)
            r = cap.add_run(caption)
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            r.font.name = 'Calibri'
            return True
        except Exception as e:
            print(f'  img err {path}: {e}')
    return False

# ── Helper: add table ─────────────────────────────────────────────────────────
def add_table_docx(doc, caption, headers, rows):
    if not headers: return
    # Caption
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(10)
    cap.paragraph_format.space_after  = Pt(4)
    r = cap.add_run(caption)
    r.bold = True; r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x0a, 0x16, 0x28)
    r.font.name = 'Calibri'

    ncols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hc = tbl.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hc[i], '0a1628')
        p = hc[i].paragraphs[0]
        p.clear()
        run = p.add_run(str(h))
        run.bold = True; run.font.size = Pt(9)
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)

    # Data rows
    for ri, row in enumerate(rows):
        rc = tbl.rows[ri + 1].cells
        bg = 'f2f5fa' if ri % 2 == 0 else 'ffffff'
        for ci, val in enumerate(list(row)[:ncols]):
            set_cell_shading(rc[ci], bg)
            p = rc[ci].paragraphs[0]
            p.clear()
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)

    doc.add_paragraph()

# ── Helper: case study ────────────────────────────────────────────────────────
def add_case_study(doc, title, body):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set column widths (gold stripe | content)
    tbl.columns[0].width = Inches(0.18)
    tbl.columns[1].width = Inches(6.32)

    gold_cell    = tbl.cell(0, 0)
    content_cell = tbl.cell(0, 1)

    set_cell_shading(gold_cell,    'c9a84c')
    set_cell_shading(content_cell, 'f2f5fa')

    # Remove all paragraphs from gold cell
    for p in gold_cell.paragraphs:
        p.clear()

    # Content: label
    content_cell.paragraphs[0].clear()
    lbl = content_cell.paragraphs[0]
    lbl.paragraph_format.space_before = Pt(8)
    lbl.paragraph_format.space_after  = Pt(2)
    r = lbl.add_run('CASE STUDY')
    r.bold = True; r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0xc9, 0xa8, 0x4c)
    r.font.name = 'Calibri'

    # Title
    ttl = content_cell.add_paragraph(title)
    ttl.paragraph_format.space_before = Pt(0)
    ttl.paragraph_format.space_after  = Pt(6)
    r2 = ttl.runs[0]
    r2.bold = True; r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x0a, 0x16, 0x28)
    r2.font.name = 'Calibri'

    # Body
    bod = content_cell.add_paragraph(body)
    bod.paragraph_format.space_before = Pt(0)
    bod.paragraph_format.space_after  = Pt(10)
    r3 = bod.runs[0]
    r3.font.size = Pt(10)
    r3.font.name = 'Times New Roman'
    r3.font.color.rgb = RGBColor(0x4a, 0x4a, 0x4a)

    doc.add_paragraph()

# ── Helper: figure/code box ───────────────────────────────────────────────────
def add_figure_box(doc, title, content):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, '111827')

    cell.paragraphs[0].clear()

    # Title in gold
    tp = cell.paragraphs[0]
    tp.paragraph_format.space_before = Pt(8)
    tp.paragraph_format.space_after  = Pt(4)
    r = tp.add_run(title)
    r.bold = True; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xc9, 0xa8, 0x4c)
    r.font.name = 'Calibri'

    # Content lines in courier
    for line in content.split('\n'):
        line = line.strip()
        if line:
            p = cell.add_paragraph(line)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(1)
            r2 = p.runs[0]
            r2.font.size = Pt(9)
            r2.font.name = 'Courier New'
            r2.font.color.rgb = RGBColor(0xe2, 0xe8, 0xf0)
    doc.add_paragraph()

# ── Title page ────────────────────────────────────────────────────────────────
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(80)
tp.paragraph_format.space_after  = Pt(6)
r = tp.add_run('Python Programming Fundamentals')
r.font.name = 'Calibri'; r.font.size = Pt(28); r.bold = True
r.font.color.rgb = RGBColor(0x0a, 0x16, 0x28)

sp = doc.add_paragraph('A Comprehensive Guide')
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp.runs[0].font.name = 'Calibri'; sp.runs[0].font.size = Pt(15)
sp.runs[0].font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

doc.add_paragraph()
ap = doc.add_paragraph('Bernard Baah')
ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
ap.runs[0].font.name = 'Calibri'; ap.runs[0].font.size = Pt(13); ap.runs[0].bold = True

fp = doc.add_paragraph('A Comprehensive Guide to Python Programming')
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.runs[0].font.name = 'Calibri'; fp.runs[0].font.size = Pt(11)
fp.runs[0].font.color.rgb = RGBColor(0xc9, 0xa8, 0x4c)

doc.add_page_break()

# ── Table of Contents ──────────────────────────────────────────────────────────
toc_h = doc.add_heading('Table of Contents', level=1)
toc_h.runs[0].font.color.rgb = RGBColor(0x0a, 0x16, 0x28)
toc_h.runs[0].font.name = 'Calibri'

CH_TITLES = {
    1:'Introduction to Python', 2:'Getting Started',
    3:'Variables and Data Types', 4:'Control Structures',
    5:'Functions', 6:'Modules and Libraries',
    7:'File Handling', 8:'Exception Handling',
    9:'Data Structures', 10:'Object-Oriented Programming: Basics',
    11:'Object-Oriented Programming: Advanced Concepts',
    12:'Error Handling and Debugging',
    13:'Working with Files and Directories (Part 2)',
    14:'Introduction to Testing',
    15:'Introduction to Modules and Packages',
    16:'Advanced Modules and Packages',
    17:'Working with External APIs',
    18:'Introduction to Web Development with Flask',
    19:'Intermediate Flask Development',
    20:'Introduction to Data Visualization with Matplotlib',
    21:'Advanced Data Visualization with Seaborn',
    22:'Introduction to Machine Learning with scikit-learn',
    23:'Web Scraping with Python',
    24:'Python Automation',
    25:'Conclusion and Next Steps',
}

PARTS = {1:'Part I: Python Foundations', 6:'Part II: Intermediate Python',
         13:'Part III: Advanced Topics', 17:'Part IV: Web, Data & Machine Learning'}

for n in range(1, 26):
    if n in PARTS:
        pp = doc.add_paragraph()
        pp.paragraph_format.space_before = Pt(8)
        r = pp.add_run(PARTS[n])
        r.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xc9, 0xa8, 0x4c)
        r.font.name = 'Calibri'
    title = ch_data[n]['title'] if n in ch_data else CH_TITLES.get(n, f'Chapter {n}')
    ep = doc.add_paragraph(f'Chapter {n}  —  {title}')
    ep.runs[0].font.size = Pt(11)
    ep.runs[0].font.name = 'Times New Roman'

doc.add_page_break()

# ── Chapters ───────────────────────────────────────────────────────────────────
def split_text(text, max_chars=1600):
    paras = [p.strip() for p in text.split('\n') if p.strip()]
    chunks, cur, clen = [], [], 0
    for p in paras:
        if clen + len(p) > max_chars and cur:
            chunks.append(cur); cur, clen = [], 0
        cur.append(p); clen += len(p)
    if cur: chunks.append(cur)
    return chunks

for ch_num in range(1, 26):
    ch    = ch_data.get(ch_num, {})
    title = ch.get('title', CH_TITLES.get(ch_num, f'Chapter {ch_num}'))
    images = ch.get('images', [])
    tables = ch.get('tables', [])
    css    = ch.get('case_studies', [])
    figs   = ch.get('text_figures', [])
    orig   = sanitize(ch_text.get(str(ch_num), ''))

    print(f'  Ch {ch_num}: {title}')

    add_chapter_banner(doc, ch_num, title)

    chunks = split_text(orig, max_chars=1600)
    img_i = tbl_i = 0
    first_para = True

    for ci, chunk in enumerate(chunks):
        for line in chunk:
            line = line.strip()
            if not line: continue
            if re.match(r'^(CHAPTER\s+\d+)', line, re.I): continue
            elif re.match(r'^\d+\.\d+\s+\S', line) or (len(line) < 65 and line.upper() == line and len(line) > 5):
                add_h2(doc, line)
                first_para = True
            elif line.startswith(('▪','•','-','–')):
                bp = doc.add_paragraph(style='List Bullet')
                bp.paragraph_format.space_after = Pt(3)
                r = bp.add_run(line.lstrip('▪•-– ').strip())
                r.font.size = Pt(11); r.font.name = 'Times New Roman'
            else:
                add_body(doc, line, first=first_para)
                first_para = False

        # Image every 2 chunks
        if img_i < len(images) and ci % 2 == 1:
            img  = images[img_i]
            path = img.get('path', img) if isinstance(img, dict) else img
            cap  = img.get('caption', f'Figure {ch_num}.{img_i+1}') if isinstance(img, dict) else f'Figure {ch_num}.{img_i+1}'
            if add_image_safe(doc, path, cap): img_i += 1

        # Table every 3 chunks
        if tbl_i < len(tables) and ci % 3 == 2:
            t = tables[tbl_i]
            try:
                add_table_docx(doc, t.get('caption',''), t.get('headers',[]), t.get('rows',[]))
                tbl_i += 1
            except Exception as e:
                print(f'    tbl err: {e}')

    while img_i < len(images):
        img  = images[img_i]
        path = img.get('path', img) if isinstance(img, dict) else img
        cap  = img.get('caption', f'Figure {ch_num}.{img_i+1}') if isinstance(img, dict) else f'Figure {ch_num}.{img_i+1}'
        if add_image_safe(doc, path, cap): img_i += 1
        else: img_i += 1

    while tbl_i < len(tables):
        t = tables[tbl_i]
        try: add_table_docx(doc, t.get('caption',''), t.get('headers',[]), t.get('rows',[]))
        except Exception as e: print(f'    tbl err: {e}')
        tbl_i += 1

    for fig in figs:
        add_figure_box(doc, fig.get('title','Figure'), fig.get('content',''))

    for cs in css:
        add_case_study(doc, cs.get('title','Case Study'), cs.get('body',''))

    doc.add_page_break()

OUT = 'book_data/Python_Fundamentals_Interior.docx'
doc.save(OUT)
sz = os.path.getsize(OUT) / 1e6
print(f'\n✓ DOCX → {OUT}  ({sz:.1f} MB)')
